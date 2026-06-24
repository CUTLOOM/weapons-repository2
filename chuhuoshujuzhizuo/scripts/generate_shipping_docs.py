#!/usr/bin/env python3
import sys, os, shutil, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from datetime import datetime
import openpyxl
from docx import Document
def extract_pi_data(pi_path):
    wb = openpyxl.load_workbook(pi_path, data_only=True)
    ws = wb.active
    def g(r, c):
        v = ws.cell(r, c).value
        return str(v).strip() if v else ''
    data = {
        'container': g(5,2), 'seal': g(6,2), 'bl': g(7,2),
        'destination': g(8,2), 'vessel': g(9,2),
        'total_ctn': ws.cell(10,2).value or 0,
        'total_gw': ws.cell(11,2).value or 0,
        'total_nw': ws.cell(12,2).value or 0,
        'total_cbm': ws.cell(13,2).value or 0,
        'customer': g(15,2),
    }
    data['inv_no'] = ws.cell(16,13).value or ws.cell(17,13).value or ''
    data['order_no'] = ws.cell(17,13).value or ws.cell(18,13).value or ''
    for r in range(ws.max_row+1):
        v = ws.cell(r,1).value
        if v and 'Spare' in str(v):
            data['spare_ctn'] = ws.cell(r,10).value or 0
            data['spare_gw'] = ws.cell(r,18).value or 0
            data['spare_cbm'] = ws.cell(r,12).value or 0
        if v and str(v).upper() == 'TOTAL':
            data['total_pcs'] = ws.cell(r,6).value or 0
            data['total_amount'] = ws.cell(r,8).value or 0
    wb.close()
    return data

def sc(cell, text):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    cell.paragraphs[0].text = text

def run(pi_path, tpl_dir, ics2_tpl, out_dir, gen_date):
    os.makedirs(out_dir, exist_ok=True)
    d = extract_pi_data(pi_path)
    chair_ctn = int(d['total_ctn']) - int(d.get('spare_ctn', 0))
    chair_gw = float(d['total_gw']) - float(d.get('spare_gw', 0))
    chair_cbm = float(d['total_cbm']) - float(d.get('spare_cbm', 0))
    avg_p = float(d['total_amount']) / int(d['total_pcs']) if int(d['total_pcs']) else 0
    print('Order:', d['order_no'], '| BL:', d['bl'], '| CTN:', d['total_ctn'])

    f1 = os.path.join(tpl_dir, '1. \u5ba2\u6237\u6837\u5355.xlsx')
    if os.path.exists(f1):
        shutil.copy2(f1, os.path.join(out_dir, '1. \u5ba2\u6237\u6837\u5355.xlsx'))
        wb = openpyxl.load_workbook(os.path.join(out_dir, '1. \u5ba2\u6237\u6837\u5355.xlsx'))
        ws = wb.active
        ws['F6'] = d['bl']; ws['A16'] = d['customer']; ws['A22'] = d['vessel']
        ws['C28'] = str(d['total_ctn']) + 'CARTONS'
        ws['D28'] = 'DINING CHAIR\nSpare parts\nOrder No: ' + d['order_no']
        ws['F28'] = d['total_gw']; ws['G28'] = d['total_cbm']
        ws['A37'] = d['container']; ws['C37'] = d['seal']; ws['D37'] = '40HQ'
        ws['E37'] = d['total_ctn']; ws['F37'] = d['total_gw']; ws['G37'] = d['total_cbm']
        ws['J37'] = chr(39184) + chr(26885)
        wb.save(os.path.join(out_dir, '1. \u5ba2\u6237\u6837\u5355.xlsx'))
        print('1. OK')

    shutil.copy2(ics2_tpl, os.path.join(out_dir, '2. ICS2.xlsx'))
    wb = openpyxl.load_workbook(os.path.join(out_dir, '2. ICS2.xlsx'))
    w1, w2 = wb.worksheets[0], wb.worksheets[1]
    w1['E6'] = 'V.031E'; w1['S6'] = 1; w1['U6'] = d['total_gw']
    w1['Y6'] = 'TIANJIN DE YUAN INTERNATIONAL TRADE CO.,LTD'
    w1['AF6'] = d['customer']

    w2['B6'] = 1; w2['C6'] = d['container']; w2['D6'] = d['seal']; w2['E6'] = '40HQ'
    w2['F6'] = 'DINING CHAIR'; w2['G6'] = 940171; w2['H6'] = chair_ctn
    w2['J6'] = chair_gw; w2['K6'] = chair_cbm; w2['L6'] = 'N/M'
    w2['B7'] = 2; w2['C7'] = d['container']; w2['D7'] = d['seal']; w2['E7'] = '40HQ'
    w2['F7'] = 'SPARE PARTS(Hardware+Feet)'; w2['G7'] = 940171
    w2['H7'] = d.get('spare_ctn', 1); w2['J7'] = d.get('spare_gw', 2); w2['K7'] = d.get('spare_cbm', 0.014)
    wb.save(os.path.join(out_dir, '2. ICS2.xlsx'))
    print('2. OK')

    pk = os.path.join(tpl_dir, '3. \u7bb1\u5355.xlsx')
    if os.path.exists(pk):
        shutil.copy2(pk, os.path.join(out_dir, '3. \u7bb1\u5355.xlsx'))
        wb = openpyxl.load_workbook(os.path.join(out_dir, '3. \u7bb1\u5355.xlsx'))
        ws = wb.active
        ws['A10'] = d['customer']; ws['D15'] = d['inv_no']; ws['E15'] = gen_date
        ws['A19'] = str(d['total_pcs']) + 'PCS'
        ws['B19'] = 'DINING CHAIR ' + str(d['total_ctn']) + 'CTNS'
        ws['C19'] = d['total_gw']; ws['D19'] = d['total_nw']; ws['E19'] = d['total_cbm']
        wb.save(os.path.join(out_dir, '3. \u7bb1\u5355.xlsx'))
        print('3. OK')

    inv_name = '4. \u53d1\u7968' + d['inv_no'] + '.xlsx'
    inv_src = os.path.join(tpl_dir, '4. \u53d1\u7968DY-H-260403035.xlsx')
    if os.path.exists(inv_src):
        shutil.copy2(inv_src, os.path.join(out_dir, inv_name))
        wb = openpyxl.load_workbook(os.path.join(out_dir, inv_name))
        ws = wb.active
        ws['A8'] = d['customer']; ws['C12'] = d['inv_no']; ws['F12'] = gen_date
        ws['C14'] = d['inv_no']
        ws['D23'] = d['total_pcs']; ws['E23'] = round(avg_p, 3); ws['G23'] = d['total_amount']
        ws['D25'] = d['total_pcs']; ws['G25'] = 'USD' + str(round(d['total_amount'], 2))
        wb.save(os.path.join(out_dir, inv_name))
        print('4. OK')

    cd_src = os.path.join(tpl_dir, '5. \u62a5\u5173\u5355.docx')
    if os.path.exists(cd_src):
        shutil.copy2(cd_src, os.path.join(out_dir, '5. \u62a5\u5173\u5355.docx'))
        doc = Document(os.path.join(out_dir, '5. \u62a5\u5173\u5355.docx'))
        t = doc.tables[0]
        sc(t.rows[1].cells[10], chr(25552)+chr(36816)+chr(21333)+chr(21495)+'\n'+d['bl'])
        sc(t.rows[1].cells[4], chr(36816)+chr(36755)+chr(24037)+chr(20855)+chr(21517)+chr(31216))
        sc(t.rows[5].cells[0], chr(21512)+chr(21516)+chr(21327)+chr(35758)+chr(21495)+'\n'+d['inv_no'])
        sc(t.rows[5].cells[1], chr(20214)+chr(25968)+'\n'+str(d['total_ctn']))
        sc(t.rows[5].cells[6], chr(27611)+chr(37325)+'\n'+str(d['total_gw']))
        sc(t.rows[5].cells[12], chr(20928)+chr(37325)+'\n'+str(d['total_nw']))
        sc(t.rows[6].cells[0], chr(38598)+chr(35013)+chr(31665)+chr(21495)+'\n'+d['container'])
        gl = '9401719000      '+chr(39184)+chr(26885)+'      '+str(d['total_pcs'])+chr(20214)+'/'+str(d['total_nw'])+chr(21315)+chr(20811)+'      %.3f     %.2f' % (avg_p, d['total_amount'])
        sc(t.rows[10].cells[0], gl)
        doc.save(os.path.join(out_dir, '5. \u62a5\u5173\u5355.docx'))
        print('5. OK')

    print('Output:', out_dir)

if __name__ == '__main__':
    import sys
    pi_path = sys.argv[1]
    sd = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    d = extract_pi_data(pi_path)
    out_dir = os.path.join(os.path.dirname(os.path.dirname(pi_path)), chr(20986)+chr(36135)+chr(25991)+chr(20214)+d['bl'])
    run(pi_path, os.path.dirname(pi_path), os.path.join(sd, 'assets', 'ICS2_Template.xlsx'), out_dir, datetime.now())
